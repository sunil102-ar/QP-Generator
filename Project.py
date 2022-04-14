import streamlit as st
import pandas as pd
import seaborn as sns
import plotly.express as px
import altair as alt
import numpy as np
from dataclasses import dataclass
from math import sqrt
from sklearn.metrics import mean_squared_error
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
import collections
import os
import jinja2
from docx import Document
from docx.shared import Inches
from PIL import Image
from jinja2 import Environment, PackageLoader, select_autoescape, FileSystemLoader
from datetime import date
import streamlit as st
from streamlit.components.v1 import iframe

# set the style for seaborn
sns.set_style('darkgrid')
import streamlit.components.v1 as components






sidebar = st.container()

with sidebar:   
    #st.sidebar.button("Add Questions")
    st.title("🎓 Welcome to the Exam Paper Generator")

from PIL import Image
st.sidebar.image('LOGO_GL.jpg', use_column_width=True)


data = [['Section A (8 Marks)', "Descriptive"], ['Section B( 17 Marks)', "Coding"], ['Section C (25 Marks)', "Coding"]]
 
# Create the pandas DataFrame
data_1 = pd.DataFrame(data, columns = ['Mid Exam Paper', 'Question Format'])

#st.table(data_1)
data_f = [['Section A (20 Marks)', "Coding"], ['Section B( 20 Marks)', "Coding"], ['Section C (30 Marks)', "Coding"]]
 
# Create the pandas DataFrame
data_2 = pd.DataFrame(data_f, columns = ['Final Exam Paper', 'Question Format'])


left, right = st.columns(2)




df = pd.read_csv("data_sheet2.csv",encoding='utf8')

algorithm = st.sidebar.selectbox(
     'Select the Exam',
     ['ITP_Mid_Exam', 'ITP_Final_Exam','NPV_Mid_Exam','NPV_Final_Exam','EDA_Mid_Exam','EDA_Final_Exam','STAT_Mid_Exam','STAT_Final_Exam'])


if algorithm == 'ITP_Mid_Exam':
    data = df[df['Assessment_Subject']=='ITP']
    data = data[data["Assessment_Category"]=="Mid Exam"]
    st.write("     Mid Exam Paper Format:")
    st.table(data_1)
    df_8_Easy = data[(data["MARKS"]==8.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_17_Medium = data[(data["MARKS"]==17.0) & (data["COMPLEXITY LEVEL"]== "Medium")]  
    df_25_Medium = data[(data["MARKS"]==25.0) & (data["COMPLEXITY LEVEL"]== "Medium")]
    df8_E = df_8_Easy.sample(n = 1) 
    df17_M = df_17_Medium.sample(n = 1) 
    df25_M = df_25_Medium.sample(n = 1) 
    Mid_Exam = [df8_E, df17_M,df25_M]
    Mid_Exam_ITP = pd.concat(Mid_Exam)
    ITP_50_MARKS = Mid_Exam_ITP[["Questions","MARKS","Solutions"]]
    #ITP_50_MARKS["MARKS"] = ITP_50_MARKS["MARKS"].astype("str")
    ITP_50_MARKS
        
    def generate_paper(ITP_50_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Mid_Exam: 50 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 2 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,ITP_50_MARKS.shape[0]):
            for j in range(0,ITP_50_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (8 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(ITP_50_MARKS.iloc[i,j])+"\t("+str(ITP_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(ITP_50_MARKS.iloc[i,j+2]))
                    elif(ITP_50_MARKS.iloc[i,j+1]=='8'):
                        document.add_paragraph(str(i+1)+") "+ str(ITP_50_MARKS.iloc[i,j])+"\t("+str(ITP_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(ITP_50_MARKS.iloc[i,j+2]))
                    elif(i==1):  
                        document.add_heading("                                           Section B (17 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(ITP_50_MARKS.iloc[i,j])+"\t("+str(ITP_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_50_MARKS.iloc[i,j+2]))
                    elif(ITP_50_MARKS.iloc[i,j+1]=='17'):
                        document.add_paragraph(str(i+1)+") "+ str(ITP_50_MARKS.iloc[i,j])+"\t("+str(ITP_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(ITP_50_MARKS.iloc[i,j+2]))
                    elif(i==2):
                        document.add_heading("                                           Section C (25 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(ITP_50_MARKS.iloc[i,j])+"\t("+str(ITP_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_50_MARKS.iloc[i,j+2]))
                    elif(ITP_50_MARKS.iloc[i,j+1]=='25'):
                        document.add_paragraph(str(i+1)+") "+ str(ITP_50_MARKS.iloc[i,j])+"\t("+str(ITP_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_50_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="ITP_MID_EXAM_SET.docx",
                 mime="application/octet-stream"
               )
            st.balloons()
    
    if st.button("Generate ITP Mid Exam Paper"):
        if algorithm=='ITP_Mid_Exam':
            generate_paper(ITP_50_MARKS,"ITP_MID_EXAM ")
        elif algorithm=='NPV':
            generate_paper(ITP_50_MARKS,"NPV ")
        elif algorithm=='EDA':
            generate_paper(ITP_50_MARKS,"EDA ")
        else:
            generate_paper(ITP_50_MARKS,"STAT ")


elif algorithm == 'ITP_Final_Exam':
    data = df[df['Assessment_Subject']=='ITP']
    data = data[data["Assessment_Category"]=="Final Exam"]
    st.write("     Final Exam Paper Format:")
    st.table(data_2)
    df_5_Medium = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Medium")].sample(n = 2) 
    df_5_Easy = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_10_Medium = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df_10_Hard = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Hard")]
    df_15_Hard = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Hard")] 
    df_15_Medium = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df5_M = df_5_Medium.sample(n = 2) 
    df5_E = df_5_Easy.sample(n = 2) 
    df10_M = df_10_Medium.sample(n = 1) 
    df10_H = df_10_Hard.sample(n = 1) 
    df15_H = df_15_Hard.sample(n = 1) 
    df15_M = df_15_Medium.sample(n = 1) 
    Final_Exam = [df5_M, df5_E, df10_M,df10_H,df15_H,df15_M]
    Final_Exam_ITP = pd.concat(Final_Exam)
    ITP_70_MARKS = Final_Exam_ITP[["Questions","MARKS","Solutions"]]
    ITP_70_MARKS["MARKS"] = ITP_70_MARKS["MARKS"].astype("str")
    ITP_70_MARKS
        
    def generate_paper(ITP_70_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Final_Exam: 70 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 3 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,ITP_70_MARKS.shape[0]):
            for j in range(0,ITP_70_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(ITP_70_MARKS.iloc[i,j])+"\t("+str(ITP_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n"+"\n------:Solution:------\n"+str(ITP_70_MARKS.iloc[i,j+2]))
                    elif(ITP_70_MARKS.iloc[i,j+1]=='5'):
                        document.add_paragraph(str(i+1)+") "+ str(ITP_70_MARKS.iloc[i,j])+"\t("+str(ITP_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_70_MARKS.iloc[i,j+2]))
                    elif(i==4):  
                        document.add_heading("                                           Section B (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(ITP_70_MARKS.iloc[i,j])+"\t("+str(ITP_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_70_MARKS.iloc[i,j+2]))
                    elif(ITP_70_MARKS.iloc[i,j+1]=='10'):
                        document.add_paragraph(str(i+1)+") "+ str(ITP_70_MARKS.iloc[i,j])+"\t("+str(ITP_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_70_MARKS.iloc[i,j+2]))
                    elif(i==6):
                        document.add_heading("                                           Section C (30 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(ITP_70_MARKS.iloc[i,j])+"\t("+str(ITP_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_70_MARKS.iloc[i,j+2]))
                    elif(ITP_70_MARKS.iloc[i,j+1]=='15'):
                        document.add_paragraph(str(i+1)+") "+ str(ITP_70_MARKS.iloc[i,j])+"\t("+str(ITP_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(ITP_70_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="ITP_Set_EXAM_SET.docx",
                 mime="application/octet-stream"
               )


    if st.button("Generate ITP Final Exam Paper"):
        if algorithm=='ITP_Final_Exam':
            generate_paper(ITP_70_MARKS,"ITP_FINAL_EXAM ")
            
        elif algorithm=='NPV':
            generate_paper(ITP_70_MARKS,"NPV ")
        elif algorithm=='EDA':
            generate_paper(ITP_70_MARKS,"EDA ")
        else:
            generate_paper(ITP_70_MARKS,"STAT ")


    
elif algorithm == 'NPV_Mid_Exam':
    data = df[df['Assessment_Subject']=='NPV']
    data = data[data["Assessment_Category"]=="Mid Exam"]
    st.write("     Mid Exam Paper Format:")
    st.table(data_1)
    st.dataframe(data.head(20))
    st.write(data.shape)
    df_8_Easy = data[(data["MARKS"]==8.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_17_Medium = data[(data["MARKS"]==17.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df_25_Hard = data[(data["MARKS"]==25.0) & (data["COMPLEXITY LEVEL"]== "Hard")] 
    df8_E = df_8_Easy.sample(n = 1) 
    df17_M = df_17_Medium.sample(n = 1) 
    df25_H = df_25_Hard.sample(n = 1) 
    Mid_Exam = [df8_E,df17_M,df25_H]
    Mid_Exam_NPV = pd.concat(Mid_Exam)
    NPV_50_MARKS = Mid_Exam_NPV[["Questions","MARKS","Solutions"]]
    NPV_50_MARKS["MARKS"] = NPV_50_MARKS["MARKS"].astype("str")
    NPV_50_MARKS
        
    def generate_paper(NPV_50_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Mid_Exam: 50 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 2 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,NPV_50_MARKS.shape[0]):
            for j in range(0,NPV_50_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (8 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(NPV_50_MARKS.iloc[i,j])+"\t("+str(NPV_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(NPV_50_MARKS.iloc[i,j+2]))
                    elif(NPV_50_MARKS.iloc[i,j+1]=='8'):
                        document.add_paragraph(str(i+1)+") "+ str(NPV_50_MARKS.iloc[i,j])+"\t("+str(NPV_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(NPV_50_MARKS.iloc[i,j+2]))
                    elif(i==1):  
                        document.add_heading("                                           Section B (17 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(NPV_50_MARKS.iloc[i,j])+"\t("+str(NPV_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_50_MARKS.iloc[i,j+2]))
                    elif(NPV_50_MARKS.iloc[i,j+1]=='17'):
                        document.add_paragraph(str(i+1)+") "+ str(NPV_50_MARKS.iloc[i,j])+"\t("+str(NPV_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(NPV_50_MARKS.iloc[i,j+2]))
                    elif(i==2):
                        document.add_heading("                                           Section C (25 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(NPV_50_MARKS.iloc[i,j])+"\t("+str(NPV_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_50_MARKS.iloc[i,j+2]))
                    elif(NPV_50_MARKS.iloc[i,j+1]=='25'):
                        document.add_paragraph(str(i+1)+") "+ str(NPV_50_MARKS.iloc[i,j])+"\t("+str(NPV_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_50_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="NPV_MID_EXAM_SET.docx",
                 mime="application/octet-stream"
               )
            st.balloons()
    
    if st.button("Generate NPV Mid Exam Paper"):
        if algorithm=='NPV_Mid_Exam':
            generate_paper(NPV_50_MARKS,"NPV_MID_EXAM ")
        elif algorithm=='NPV':
            generate_paper(NPV_50_MARKS,"NPV ")
        elif algorithm=='EDA':
            generate_paper(NPV_50_MARKS,"EDA ")
        else:
            generate_paper(NPV_50_MARKS,"STAT ")


elif algorithm == 'NPV_Final_Exam':
    data = df[df['Assessment_Subject']=='NPV']
    data = data[data["Assessment_Category"]=="Final Exam"]
    st.write("     Final Exam Paper Formate:")
    st.table(data_2)
    df_5_Medium = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Medium")].sample(n = 2) 
    df_5_Easy = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_10_Medium = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df_10_Hard = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Hard")]
    df_15_Hard = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Hard")] 
    df_15_Medium = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df5_M = df_5_Medium.sample(n = 2) 
    df5_E = df_5_Easy.sample(n = 2) 
    df10_M = df_10_Medium.sample(n = 1) 
    df10_H = df_10_Hard.sample(n = 1) 
    df15_H = df_15_Hard.sample(n = 1) 
    df15_M = df_15_Medium.sample(n = 1) 
    Final_Exam = [df5_M, df5_E, df10_M,df10_H,df15_H,df15_M]
    Final_Exam_ITP = pd.concat(Final_Exam)
    NPV_70_MARKS = Final_Exam_ITP[["Questions","MARKS","Solutions"]]
    NPV_70_MARKS["MARKS"] = NPV_70_MARKS["MARKS"].astype("str")
    NPV_70_MARKS
        
    def generate_paper(NPV_70_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Final_Exam: 70 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 3 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,NPV_70_MARKS.shape[0]):
            for j in range(0,NPV_70_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(NPV_70_MARKS.iloc[i,j])+"\t("+str(NPV_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n"+"\n------:Solution:------\n"+str(NPV_70_MARKS.iloc[i,j+2]))
                    elif(NPV_70_MARKS.iloc[i,j+1]=='5'):
                        document.add_paragraph(str(i+1)+") "+ str(NPV_70_MARKS.iloc[i,j])+"\t("+str(NPV_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_70_MARKS.iloc[i,j+2]))
                    elif(i==4):  
                        document.add_heading("                                           Section B (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(NPV_70_MARKS.iloc[i,j])+"\t("+str(NPV_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_70_MARKS.iloc[i,j+2]))
                    elif(NPV_70_MARKS.iloc[i,j+1]=='10'):
                        document.add_paragraph(str(i+1)+") "+ str(NPV_70_MARKS.iloc[i,j])+"\t("+str(NPV_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_70_MARKS.iloc[i,j+2]))
                    elif(i==6):
                        document.add_heading("                                           Section C (30 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(NPV_70_MARKS.iloc[i,j])+"\t("+str(NPV_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_70_MARKS.iloc[i,j+2]))
                    elif(NPV_70_MARKS.iloc[i,j+1]=='15'):
                        document.add_paragraph(str(i+1)+") "+ str(NPV_70_MARKS.iloc[i,j])+"\t("+str(NPV_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(NPV_70_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="NPV_Set_EXAM_SET.docx",
                 mime="application/octet-stream"
               )


    if st.button("Generate NPV Final Exam Paper"):
        if algorithm=='NPV_Final_Exam':
            generate_paper(NPV_70_MARKS,"NPV_FINAL_EXAM ")
            
        elif algorithm=='NPV':
            generate_paper(NPV_70_MARKS,"NPV ")
        elif algorithm=='EDA':
            generate_paper(NPV_70_MARKS,"EDA ")
        else:
            generate_paper(NPV_70_MARKS,"STAT ")

    
      
# elif algorithm == 'EDA':
#     data= df[df['Assessment_Subject']=='EDA']
#     st.dataframe(data.head())
#     st.write(data.shape)  


elif algorithm == 'EDA_Mid_Exam':
    data = df[df['Assessment_Subject']=='EDA']
    data = data[data["Assessment_Category"]=="Mid Exam"]
    st.write("     Mid Exam Paper Format:")
    st.table(data_1)
    df_8_Easy = data[(data["MARKS"]==8.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_17_Medium = data[(data["MARKS"]==17.0) & (data["COMPLEXITY LEVEL"]== "Medium")]  
    df_25_Medium = data[(data["MARKS"]==25.0) & (data["COMPLEXITY LEVEL"]== "Medium")]
    df8_E = df_8_Easy.sample(n = 1) 
    df17_M = df_17_Medium.sample(n = 1) 
    df25_M = df_25_Medium.sample(n = 1) 
    Mid_Exam = [df8_E, df17_M,df25_M]
    Mid_Exam_EDA = pd.concat(Mid_Exam)
    EDA_50_MARKS = Mid_Exam_EDA[["Questions","MARKS","Solutions"]]
    EDA_50_MARKS["MARKS"] = EDA_50_MARKS["MARKS"].astype("str")
    EDA_50_MARKS
        
    def generate_paper(EDA_50_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Mid_Exam: 50 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 2 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,EDA_50_MARKS.shape[0]):
            for j in range(0,EDA_50_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (8 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(EDA_50_MARKS.iloc[i,j])+"\t("+str(EDA_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(EDA_50_MARKS.iloc[i,j+2]))
                    elif(EDA_50_MARKS.iloc[i,j+1]=='8'):
                        document.add_paragraph(str(i+1)+") "+ str(EDA_50_MARKS.iloc[i,j])+"\t("+str(EDA_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(EDA_50_MARKS.iloc[i,j+2]))
                    elif(i==1):  
                        document.add_heading("                                           Section B (17 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(EDA_50_MARKS.iloc[i,j])+"\t("+str(EDA_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_50_MARKS.iloc[i,j+2]))
                    elif(EDA_50_MARKS.iloc[i,j+1]=='17'):
                        document.add_paragraph(str(i+1)+") "+ str(EDA_50_MARKS.iloc[i,j])+"\t("+str(EDA_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(EDA_50_MARKS.iloc[i,j+2]))
                    elif(i==2):
                        document.add_heading("                                           Section C (25 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(EDA_50_MARKS.iloc[i,j])+"\t("+str(EDA_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_50_MARKS.iloc[i,j+2]))
                    elif(EDA_50_MARKS.iloc[i,j+1]=='25'):
                        document.add_paragraph(str(i+1)+") "+ str(EDA_50_MARKS.iloc[i,j])+"\t("+str(EDA_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_50_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="EDA_MID_EXAM_SET.docx",
                 mime="application/octet-stream"
               )
            st.balloons()
    
    if st.button("Generate EDA Mid Exam Paper"):
        if algorithm=='EDA_Mid_Exam':
            generate_paper(EDA_50_MARKS,"EDA_MID_EXAM ")
        elif algorithm=='NPV':
            generate_paper(EDA_50_MARKS,"NPV ")
        elif algorithm=='EDA':
            generate_paper(EDA_50_MARKS,"EDA ")
        else:
            generate_paper(EDA_50_MARKS,"STAT ")


elif algorithm == 'EDA_Final_Exam':
    data = df[df['Assessment_Subject']=='EDA']
    data = data[data["Assessment_Category"]=="Final Exam"]
    st.write("     Final Exam Paper Format:")
    st.table(data_2)
    df_5_Medium = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Medium")].sample(n = 2) 
    df_5_Easy = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_10_Medium = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df_10_Hard = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Hard")]
    df_15_Hard = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Hard")] 
    df_15_Medium = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df5_M = df_5_Medium.sample(n = 2) 
    df5_E = df_5_Easy.sample(n = 2) 
    df10_M = df_10_Medium.sample(n = 1) 
    df10_H = df_10_Hard.sample(n = 1) 
    df15_H = df_15_Hard.sample(n = 1) 
    df15_M = df_15_Medium.sample(n = 1) 
    Final_Exam = [df5_M, df5_E, df10_M,df10_H,df15_H,df15_M]
    Final_Exam_EDA = pd.concat(Final_Exam)
    EDA_70_MARKS = Final_Exam_EDA[["Questions","MARKS","Solutions"]]
    EDA_70_MARKS["MARKS"] = EDA_70_MARKS["MARKS"].astype("str")
    EDA_70_MARKS
        
    def generate_paper(EDA_70_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Final_Exam: 70 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 3 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,EDA_70_MARKS.shape[0]):
            for j in range(0,EDA_70_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(EDA_70_MARKS.iloc[i,j])+"\t("+str(EDA_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n"+"\n------:Solution:------\n"+str(EDA_70_MARKS.iloc[i,j+2]))
                    elif(EDA_70_MARKS.iloc[i,j+1]=='5'):
                        document.add_paragraph(str(i+1)+") "+ str(EDA_70_MARKS.iloc[i,j])+"\t("+str(EDA_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_70_MARKS.iloc[i,j+2]))
                    elif(i==4):  
                        document.add_heading("                                           Section B (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(EDA_70_MARKS.iloc[i,j])+"\t("+str(EDA_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_70_MARKS.iloc[i,j+2]))
                    elif(EDA_70_MARKS.iloc[i,j+1]=='10'):
                        document.add_paragraph(str(i+1)+") "+ str(EDA_70_MARKS.iloc[i,j])+"\t("+str(EDA_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_70_MARKS.iloc[i,j+2]))
                    elif(i==6):
                        document.add_heading("                                           Section C (30 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(EDA_70_MARKS.iloc[i,j])+"\t("+str(EDA_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_70_MARKS.iloc[i,j+2]))
                    elif(EDA_70_MARKS.iloc[i,j+1]=='15'):
                        document.add_paragraph(str(i+1)+") "+ str(EDA_70_MARKS.iloc[i,j])+"\t("+str(EDA_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(EDA_70_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="EDA_FINAL_EXAM_SET.docx",
                 mime="application/octet-stream"
               )


    if st.button("Generate EDA Final Exam Paper"):
        if algorithm=='EDA_Final_Exam':
            generate_paper(EDA_70_MARKS,"EDA_FINAL_EXAM ")
            
        elif algorithm=='NPV':
            generate_paper(EDA_70_MARKS,"NPV ")
        elif algorithm=='EDA':
            generate_paper(EDA_70_MARKS,"EDA ")
        else:
            generate_paper(EDA_70_MARKS,"STAT ")

    
elif algorithm == 'STAT_Mid_Exam':
    data = df[df['Assessment_Subject']=='Stat']
    data = data[data["Assessment_Category"]=="Mid Exam"]
    st.write("     Mid Exam Paper Format:")
    st.table(data_1)
    df_8_Easy = data[(data["MARKS"]==8.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_17_Medium = data[(data["MARKS"]==17.0) & (data["COMPLEXITY LEVEL"]== "Medium")]  
    df_25_Medium = data[(data["MARKS"]==25.0) & (data["COMPLEXITY LEVEL"]== "Medium")]
    df8_E = df_8_Easy.sample(n = 1) 
    df17_M = df_17_Medium.sample(n = 1) 
    df25_M = df_25_Medium.sample(n = 1) 
    Mid_Exam = [df8_E, df17_M,df25_M]
    Mid_Exam_STAT = pd.concat(Mid_Exam)
    STAT_50_MARKS = Mid_Exam_STAT[["Questions","MARKS","Solutions"]]
    STAT_50_MARKS["MARKS"] = STAT_50_MARKS["MARKS"].astype("str")
    STAT_50_MARKS
        
    def generate_paper(STAT_50_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Mid_Exam: 50 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 2 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,STAT_50_MARKS.shape[0]):
            for j in range(0,STAT_50_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (8 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(STAT_50_MARKS.iloc[i,j])+"\t("+str(STAT_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(STAT_50_MARKS.iloc[i,j+2]))
                    elif(STAT_50_MARKS.iloc[i,j+1]=='8'):
                        document.add_paragraph(str(i+1)+") "+ str(STAT_50_MARKS.iloc[i,j])+"\t("+str(STAT_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(STAT_50_MARKS.iloc[i,j+2]))
                    elif(i==1):  
                        document.add_heading("                                           Section B (17 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(STAT_50_MARKS.iloc[i,j])+"\t("+str(STAT_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_50_MARKS.iloc[i,j+2]))
                    elif(STAT_50_MARKS.iloc[i,j+1]=='17'):
                        document.add_paragraph(str(i+1)+") "+ str(STAT_50_MARKS.iloc[i,j])+"\t("+str(STAT_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n "+"\n------:Solution:------\n"+str(STAT_50_MARKS.iloc[i,j+2]))
                    elif(i==2):
                        document.add_heading("                                           Section C (25 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(STAT_50_MARKS.iloc[i,j])+"\t("+str(STAT_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_50_MARKS.iloc[i,j+2]))
                    elif(STAT_50_MARKS.iloc[i,j+1]=='25'):
                        document.add_paragraph(str(i+1)+") "+ str(STAT_50_MARKS.iloc[i,j])+"\t("+str(STAT_50_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_50_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="STAT_MID_EXAM_SET.docx",
                 mime="application/octet-stream"
               )
            st.balloons()
    
    if st.button("Generate STAT Mid Exam Paper"):
        if algorithm=='STAT_Mid_Exam':
            generate_paper(STAT_50_MARKS,"STAT_MID_EXAM ")
        #elif algorithm=='NPV':
           # generate_paper(EDA_50_MARKS,"NPV ")
        #elif algorithm=='EDA':
            #generate_paper(EDA_50_MARKS,"EDA ")
        ##else:
          #  generate_paper(ITP_50_MARKS,"ITP ")


elif algorithm == 'STAT_Final_Exam':
    data = df[df['Assessment_Subject']=='Stat']
    data = data[data["Assessment_Category"]=="Final Exam"]
    st.write("     Final Exam Paper Format:")
    st.table(data_2)
    df_5_Medium = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Medium")].sample(n = 2) 
    df_5_Easy = data[(data["MARKS"]==5.0) & (data["COMPLEXITY LEVEL"]== "Easy")] 
    df_10_Medium = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df_10_Hard = data[(data["MARKS"]==10.0) & (data["COMPLEXITY LEVEL"]== "Hard")]
    df_15_Hard = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Hard")] 
    df_15_Medium = data[(data["MARKS"]==15.0) & (data["COMPLEXITY LEVEL"]== "Medium")] 
    df5_M = df_5_Medium.sample(n = 2) 
    df5_E = df_5_Easy.sample(n = 2) 
    df10_M = df_10_Medium.sample(n = 1) 
    df10_H = df_10_Hard.sample(n = 1) 
    df15_H = df_15_Hard.sample(n = 1) 
    df15_M = df_15_Medium.sample(n = 1) 
    Final_Exam = [df5_M, df5_E, df10_M,df10_H,df15_H,df15_M]
    Final_Exam_STAT = pd.concat(Final_Exam)
    STAT_70_MARKS = Final_Exam_STAT[["Questions","MARKS","Solutions"]]
    STAT_70_MARKS["MARKS"] = STAT_70_MARKS["MARKS"].astype("str")
    STAT_70_MARKS
        
    def generate_paper(STAT_70_MARKS, HEADING):
        document = Document()
        document.add_picture('LOGO_GL.jpg', width=Inches(1.5))     
        document.add_heading(HEADING,0)
        document.add_paragraph('Final_Exam: 70 MARKS', style='List Bullet')
        document.add_paragraph('Duration: 3 Hours', style='List Bullet') 
        document.add_heading('INSTRUCTIONS:', level=1)
        document.add_paragraph(
        'Candidates should answer all the questions in the same order provided in the question paper.', style='List Number')
    
        document.add_paragraph(
        'Any activity that compromises the integrity of the examination will not be permitted.', style='List Number')
    
        document.add_paragraph(
        'Students should complete the examination within the provided timeline.', style='List Number')
    
        document.add_paragraph(
        'Candidates are expected to check and ensure that the correct answer file (in. ipynb format) is uploaded in LMS.', style='List Number')
    
        for i in range(0,STAT_70_MARKS.shape[0]):
            for j in range(0,STAT_70_MARKS.shape[1]):
                if(j==0):
                    if(i==0):
                        document.add_heading("                                           Section A (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(STAT_70_MARKS.iloc[i,j])+"\t("+str(STAT_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" )\n"+"\n------:Solution:------\n"+str(STAT_70_MARKS.iloc[i,j+2]))
                    elif(STAT_70_MARKS.iloc[i,j+1]=='5'):
                        document.add_paragraph(str(i+1)+") "+ str(STAT_70_MARKS.iloc[i,j])+"\t("+str(STAT_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_70_MARKS.iloc[i,j+2]))
                    elif(i==4):  
                        document.add_heading("                                           Section B (20 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(STAT_70_MARKS.iloc[i,j])+"\t("+str(STAT_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_70_MARKS.iloc[i,j+2]))
                    elif(STAT_70_MARKS.iloc[i,j+1]=='10'):
                        document.add_paragraph(str(i+1)+") "+ str(STAT_70_MARKS.iloc[i,j])+"\t("+str(STAT_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_70_MARKS.iloc[i,j+2]))
                    elif(i==6):
                        document.add_heading("                                           Section C (30 Marks)",1)
                        document.add_paragraph(str(i+1)+") "+ str(STAT_70_MARKS.iloc[i,j])+"\t("+str(STAT_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_70_MARKS.iloc[i,j+2]))
                    elif(STAT_70_MARKS.iloc[i,j+1]=='15'):
                        document.add_paragraph(str(i+1)+") "+ str(STAT_70_MARKS.iloc[i,j])+"\t("+str(STAT_70_MARKS.iloc[i,j+1])+":- " + " "+ "Marks"+" ) \n"+"\n------:Solution:------\n"+str(STAT_70_MARKS.iloc[i,j+2])) 
                    
        document.add_heading("                                                                               ",0)      
    
        document.save(HEADING+'.docx')
        with open(HEADING+'.docx', "rb") as file:
            btn = st.download_button(
                 label="📓 Download (.iDoc)",
                 data=file,
                 file_name="STAT_FINAL_EXAM_SET.docx",
                 mime="application/octet-stream"
               )


    if st.button("Generate STAT Final Exam Paper"):
        if algorithm=='STAT_Final_Exam':
            generate_paper(STAT_70_MARKS,"STAT_FINAL_EXAM ")
            
       # elif algorithm=='NPV':
           # generate_paper(EDA_70_MARKS,"NPV ")
        #elif algorithm=='EDA':
          #  generate_paper(EDA_70_MARKS,"EDA ")
        #else:
           # generate_paper(ITP_70_MARKS,"ITP")  




st.write("Upload your Word files and convert them to Jupyter notebook files:https://alldocs.app/convert-word-docx-to-jupyter-notebook")

















